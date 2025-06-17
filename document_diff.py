#!/usr/bin/env python3

import os
import sys
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import difflib
import re

# Enhanced imports for document processing
import pdfplumber  # Better PDF handling
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
import textdistance
from pathvalidate import sanitize_filename
import Levenshtein
from jinja2 import Template

class EnhancedDocumentComparer:
    """Enhanced document comparison with multiple algorithms for web interface"""
    
    def __init__(self):
        self.supported_formats = {'.docx', '.pdf', '.txt'}
        self.algorithms = {
            'unified': 'Standard unified diff',
            'context': 'Context diff with surrounding lines',
            'levenshtein': 'Advanced Levenshtein distance',
            'jaro_winkler': 'Jaro-Winkler similarity',
            'semantic': 'Semantic similarity comparison'
        }
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from various file formats with metadata"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        metadata = {
            'file_size': file_path.stat().st_size,
            'modified_time': datetime.fromtimestamp(file_path.stat().st_mtime),
            'file_type': file_path.suffix.lower()
        }
        
        if file_path.suffix.lower() == '.pdf':
            text, pdf_meta = self._extract_pdf_text(file_path)
            metadata.update(pdf_meta)
        elif file_path.suffix.lower() == '.docx':
            text, docx_meta = self._extract_docx_text(file_path)
            metadata.update(docx_meta)
        elif file_path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            metadata['encoding'] = 'utf-8'
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        return text, metadata
    
    def _extract_pdf_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Enhanced PDF text extraction with pdfplumber"""
        text = ""
        metadata = {
            'pages': 0,
            'has_images': False,
            'fonts': set(),
            'creation_date': '',
            'author': '',
            'title': ''
        }
        
        with pdfplumber.open(str(file_path)) as pdf:
            metadata['pages'] = len(pdf.pages)
            
            # Try to get metadata
            if hasattr(pdf, 'metadata') and pdf.metadata:
                metadata['creation_date'] = pdf.metadata.get('CreationDate', '')
                metadata['author'] = pdf.metadata.get('Author', '')
                metadata['title'] = pdf.metadata.get('Title', '')
            
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"[PAGE {page_num + 1}]\n{page_text}\n\n"
                
                # Check for images
                if hasattr(page, 'images') and page.images:
                    metadata['has_images'] = True
                
                # Get font information if available
                if hasattr(page, 'chars'):
                    chars = page.chars
                    for char in chars[:100]:  # Limit to avoid performance issues
                        if 'fontname' in char:
                            metadata['fonts'].add(char['fontname'])
        
        metadata['fonts'] = list(metadata['fonts'])
        
        return text, metadata
    
    def _extract_docx_text(self, file_path: Path) -> Tuple[str, Dict]:
        """Enhanced DOCX text extraction with formatting info"""
        doc = Document(str(file_path))
        text = ""
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'has_images': False,
            'styles': set()
        }
        
        # Extract paragraph text with basic formatting info
        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name if para.style else 'Normal'
                metadata['styles'].add(style_name)
                text += f"{para.text}\n"
        
        # Extract table content
        for table in doc.tables:
            text += "\n[TABLE]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text += " | ".join(row_text) + "\n"
            text += "[/TABLE]\n\n"
        
        # Check for images (basic check)
        if hasattr(doc, 'inline_shapes'):
            metadata['has_images'] = len(doc.inline_shapes) > 0
        
        metadata['styles'] = list(metadata['styles'])
        
        return text, metadata
    
    def compare_documents(self, file1: str, file2: str, algorithm: str = 'unified') -> Dict:
        """Compare two documents using specified algorithm"""
        
        # Extract text and metadata
        text1, meta1 = self.extract_text_from_file(file1)
        text2, meta2 = self.extract_text_from_file(file2)
        
        # Perform comparison based on algorithm
        comparison_result = {
            'files': {
                'file1': {'path': file1, 'metadata': meta1},
                'file2': {'path': file2, 'metadata': meta2}
            },
            'algorithm': algorithm,
            'timestamp': datetime.now().isoformat(),
            'statistics': {},
            'differences': []
        }
        
        if algorithm == 'unified':
            comparison_result['differences'] = self._unified_diff(text1, text2)
        elif algorithm == 'context':
            comparison_result['differences'] = self._context_diff(text1, text2)
        elif algorithm == 'levenshtein':
            comparison_result['differences'] = self._levenshtein_diff(text1, text2)
        elif algorithm == 'jaro_winkler':
            comparison_result['differences'] = self._jaro_winkler_diff(text1, text2)
        elif algorithm == 'semantic':
            comparison_result['differences'] = self._semantic_diff(text1, text2)
        
        # Calculate statistics
        comparison_result['statistics'] = self._calculate_statistics(text1, text2, comparison_result['differences'])
        
        return comparison_result
    
    def _unified_diff(self, text1: str, text2: str) -> List[Dict]:
        """Enhanced unified diff with line numbers and context"""
        lines1 = text1.splitlines(keepends=True)
        lines2 = text2.splitlines(keepends=True)
        
        diff = list(difflib.unified_diff(
            lines1, lines2,
            fromfile='File 1',
            tofile='File 2',
            lineterm='',
            n=3  # More context lines
        ))
        
        differences = []
        for i, line in enumerate(diff):
            if line.startswith('@@'):
                # Parse hunk header
                match = re.match(r'@@ -(\d+),?(\d*) \+(\d+),?(\d*) @@', line)
                if match:
                    differences.append({
                        'type': 'hunk_header',
                        'line_number': i,
                        'content': line.strip(),
                        'old_start': int(match.group(1)),
                        'old_count': int(match.group(2)) if match.group(2) else 1,
                        'new_start': int(match.group(3)),
                        'new_count': int(match.group(4)) if match.group(4) else 1
                    })
            elif line.startswith('+'):
                differences.append({
                    'type': 'addition',
                    'line_number': i,
                    'content': line[1:].rstrip()
                })
            elif line.startswith('-'):
                differences.append({
                    'type': 'deletion',
                    'line_number': i,
                    'content': line[1:].rstrip()
                })
            elif not line.startswith('+++') and not line.startswith('---'):
                differences.append({
                    'type': 'context',
                    'line_number': i,
                    'content': line.rstrip()
                })
        
        return differences
    
    def _context_diff(self, text1: str, text2: str) -> List[Dict]:
        """Context diff with more surrounding lines"""
        lines1 = text1.splitlines(keepends=True)
        lines2 = text2.splitlines(keepends=True)
        
        diff = list(difflib.context_diff(
            lines1, lines2,
            fromfile='File 1',
            tofile='File 2',
            lineterm='',
            n=5  # More context
        ))
        
        differences = []
        for i, line in enumerate(diff):
            if line.startswith('*** '):
                differences.append({'type': 'file1_header', 'line_number': i, 'content': line.strip()})
            elif line.startswith('--- '):
                differences.append({'type': 'file2_header', 'line_number': i, 'content': line.strip()})
            elif line.startswith('***************'):
                differences.append({'type': 'separator', 'line_number': i, 'content': line.strip()})
            elif line.startswith('+ '):
                differences.append({'type': 'addition', 'line_number': i, 'content': line[2:].rstrip()})
            elif line.startswith('- '):
                differences.append({'type': 'deletion', 'line_number': i, 'content': line[2:].rstrip()})
            elif line.startswith('! '):
                differences.append({'type': 'change', 'line_number': i, 'content': line[2:].rstrip()})
            else:
                differences.append({'type': 'context', 'line_number': i, 'content': line.rstrip()})
        
        return differences
    
    def _levenshtein_diff(self, text1: str, text2: str) -> List[Dict]:
        """Levenshtein distance-based comparison"""
        distance = Levenshtein.distance(text1, text2)
        ratio = Levenshtein.ratio(text1, text2)
        
        # Get edit operations
        ops = Levenshtein.editops(text1, text2)
        
        differences = []
        differences.append({
            'type': 'summary',
            'content': f"Levenshtein Distance: {distance}",
            'similarity_ratio': ratio,
            'edit_operations': len(ops)
        })
        
        # Convert edit operations to readable format
        for i, (op, pos1, pos2) in enumerate(ops[:100]):  # Limit to first 100 operations
            if op == 'replace':
                differences.append({
                    'type': 'replacement',
                    'position1': pos1,
                    'position2': pos2,
                    'old_char': text1[pos1] if pos1 < len(text1) else '',
                    'new_char': text2[pos2] if pos2 < len(text2) else ''
                })
            elif op == 'insert':
                differences.append({
                    'type': 'insertion',
                    'position': pos2,
                    'char': text2[pos2] if pos2 < len(text2) else ''
                })
            elif op == 'delete':
                differences.append({
                    'type': 'deletion',
                    'position': pos1,
                    'char': text1[pos1] if pos1 < len(text1) else ''
                })
        
        return differences
    
    def _jaro_winkler_diff(self, text1: str, text2: str) -> List[Dict]:
        """Jaro-Winkler similarity comparison"""
        similarity = textdistance.jaro_winkler(text1, text2)
        
        differences = [{
            'type': 'similarity_score',
            'jaro_winkler_similarity': similarity,
            'percentage': f"{similarity * 100:.2f}%"
        }]
        
        # Add word-level comparison for more detail
        words1 = text1.split()
        words2 = text2.split()
        
        word_similarities = []
        max_len = max(len(words1), len(words2))
        
        for i in range(min(50, max_len)):  # Limit to first 50 words
            w1 = words1[i] if i < len(words1) else ""
            w2 = words2[i] if i < len(words2) else ""
            
            if w1 or w2:
                word_sim = textdistance.jaro_winkler(w1, w2)
                word_similarities.append({
                    'position': i,
                    'word1': w1,
                    'word2': w2,
                    'similarity': word_sim
                })
        
        differences.append({
            'type': 'word_similarities',
            'content': word_similarities
        })
        
        return differences
    
    def _semantic_diff(self, text1: str, text2: str) -> List[Dict]:
        """Semantic similarity using multiple algorithms"""
        algorithms = {
            'jaccard': textdistance.jaccard,
            'cosine': textdistance.cosine,
            'overlap': textdistance.overlap,
            'sorensen_dice': textdistance.sorensen_dice
        }
        
        differences = []
        
        for name, algorithm in algorithms.items():
            try:
                similarity = algorithm(text1, text2)
                differences.append({
                    'type': 'semantic_score',
                    'algorithm': name,
                    'similarity': similarity,
                    'percentage': f"{similarity * 100:.2f}%"
                })
            except Exception as e:
                differences.append({
                    'type': 'semantic_error',
                    'algorithm': name,
                    'error': str(e)
                })
        
        return differences
    
    def _calculate_statistics(self, text1: str, text2: str, differences: List[Dict]) -> Dict:
        """Calculate comprehensive statistics"""
        stats = {
            'text1': {
                'characters': len(text1),
                'words': len(text1.split()),
                'lines': len(text1.splitlines()),
                'unique_words': len(set(text1.lower().split()))
            },
            'text2': {
                'characters': len(text2),
                'words': len(text2.split()),
                'lines': len(text2.splitlines()),
                'unique_words': len(set(text2.lower().split()))
            },
            'differences': {
                'total_changes': len([d for d in differences if d.get('type') in ['addition', 'deletion', 'replacement']]),
                'additions': len([d for d in differences if d.get('type') == 'addition']),
                'deletions': len([d for d in differences if d.get('type') == 'deletion']),
                'modifications': len([d for d in differences if d.get('type') in ['replacement', 'change']])
            }
        }
        
        # Calculate similarity percentage
        total_chars = max(stats['text1']['characters'], stats['text2']['characters'])
        if total_chars > 0:
            char_diff = abs(stats['text1']['characters'] - stats['text2']['characters'])
            similarity = max(0, (total_chars - char_diff) / total_chars * 100)
            stats['similarity_percentage'] = round(similarity, 2)
        else:
            stats['similarity_percentage'] = 100.0
        
        return stats
    
    # Note: display_comparison_results method removed as it was for CLI interface
    
    def export_results(self, results: Dict, output_path: str, format: str = 'docx'):
        """Export results in various formats"""
        
        if format == 'docx':
            self._export_to_docx(results, output_path)
        elif format == 'html':
            self._export_to_html(results, output_path)
        elif format == 'json':
            self._export_to_json(results, output_path)
        elif format == 'markdown':
            self._export_to_markdown(results, output_path)
        else:
            raise ValueError(f"Unsupported export format: {format}")
    
    def _export_to_docx(self, results: Dict, output_path: str):
        """Enhanced DOCX export with better formatting"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Enhanced Document Comparison Report', 0)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Algorithm: {results['algorithm']}")
        
        # File information
        doc.add_heading('Files Compared', level=1)
        doc.add_paragraph(f"File 1: {results['files']['file1']['path']}")
        doc.add_paragraph(f"File 2: {results['files']['file2']['path']}")
        
        # Statistics
        doc.add_heading('Statistics', level=1)
        stats = results['statistics']
        
        # Create a table for statistics
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'File 1'
        hdr_cells[2].text = 'File 2'
        hdr_cells[3].text = 'Difference'
        
        # Add statistics rows
        for metric in ['characters', 'words', 'lines']:
            row_cells = table.add_row().cells
            row_cells[0].text = metric.capitalize()
            row_cells[1].text = f"{stats['text1'][metric]:,}"
            row_cells[2].text = f"{stats['text2'][metric]:,}"
            row_cells[3].text = f"{abs(stats['text1'][metric] - stats['text2'][metric]):,}"
        
        # Similarity
        doc.add_heading('Similarity Analysis', level=1)
        doc.add_paragraph(f"Overall Similarity: {stats.get('similarity_percentage', 0):.2f}%")
        
        # Changes summary
        changes = stats['differences']
        if changes['total_changes'] > 0:
            doc.add_heading('Changes Summary', level=1)
            doc.add_paragraph(f"Total Changes: {changes['total_changes']}")
            doc.add_paragraph(f"Additions: {changes['additions']}")
            doc.add_paragraph(f"Deletions: {changes['deletions']}")
            doc.add_paragraph(f"Modifications: {changes['modifications']}")
        
        # Detailed differences (limited to prevent huge files)
        doc.add_heading('Detailed Differences (First 100)', level=1)
        for i, diff in enumerate(results['differences'][:100]):
            if diff.get('type') == 'addition':
                p = doc.add_paragraph()
                run = p.add_run(f"+ {diff.get('content', '')}")
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif diff.get('type') == 'deletion':
                p = doc.add_paragraph()
                run = p.add_run(f"- {diff.get('content', '')}")
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif diff.get('type') == 'context':
                p = doc.add_paragraph()
                run = p.add_run(f"  {diff.get('content', '')}")
        
        doc.save(output_path)
        # Report saved successfully
    
    def _export_to_html(self, results: Dict, output_path: str):
        """Export to HTML with syntax highlighting"""
        html_template = """
<!DOCTYPE html>
<html>
<head>
    <title>Document Comparison Report</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        .header { background-color: #f0f0f0; padding: 15px; border-radius: 5px; }
        .stats-table { border-collapse: collapse; width: 100%; margin: 20px 0; }
        .stats-table th, .stats-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
        .stats-table th { background-color: #f2f2f2; }
        .addition { color: green; font-weight: bold; }
        .deletion { color: red; font-weight: bold; }
        .context { color: #666; }
        .similarity { font-size: 24px; text-align: center; padding: 20px; }
        .similarity.high { color: green; }
        .similarity.medium { color: orange; }
        .similarity.low { color: red; }
    </style>
</head>
<body>
    <div class="header">
        <h1>Enhanced Document Comparison Report</h1>
        <p>Generated: {{ timestamp }}</p>
        <p>Algorithm: {{ algorithm }}</p>
    </div>
    
    <h2>Files Compared</h2>
    <ul>
        <li>File 1: {{ file1_path }}</li>
        <li>File 2: {{ file2_path }}</li>
    </ul>
    
    <h2>Statistics</h2>
    <table class="stats-table">
        <tr><th>Metric</th><th>File 1</th><th>File 2</th><th>Difference</th></tr>
        <tr><td>Characters</td><td>{{ stats.text1.characters }}</td><td>{{ stats.text2.characters }}</td><td>{{ char_diff }}</td></tr>
        <tr><td>Words</td><td>{{ stats.text1.words }}</td><td>{{ stats.text2.words }}</td><td>{{ word_diff }}</td></tr>
        <tr><td>Lines</td><td>{{ stats.text1.lines }}</td><td>{{ stats.text2.lines }}</td><td>{{ line_diff }}</td></tr>
    </table>
    
    <div class="similarity {{ similarity_class }}">
        Similarity: {{ similarity }}%
    </div>
    
    <h2>Differences</h2>
    <div class="differences">
        {% for diff in differences[:100] %}
            {% if diff.type == 'addition' %}
                <div class="addition">+ {{ diff.content }}</div>
            {% elif diff.type == 'deletion' %}
                <div class="deletion">- {{ diff.content }}</div>
            {% elif diff.type == 'context' %}
                <div class="context">  {{ diff.content }}</div>
            {% endif %}
        {% endfor %}
    </div>
</body>
</html>
        """
        
        template = Template(html_template)
        stats = results['statistics']
        similarity = stats.get('similarity_percentage', 0)
        
        html_content = template.render(
            timestamp=results['timestamp'],
            algorithm=results['algorithm'],
            file1_path=results['files']['file1']['path'],
            file2_path=results['files']['file2']['path'],
            stats=stats,
            char_diff=abs(stats['text1']['characters'] - stats['text2']['characters']),
            word_diff=abs(stats['text1']['words'] - stats['text2']['words']),
            line_diff=abs(stats['text1']['lines'] - stats['text2']['lines']),
            similarity=f"{similarity:.2f}",
            similarity_class='high' if similarity > 80 else 'medium' if similarity > 50 else 'low',
            differences=results['differences']
        )
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        console.print(f"[green]HTML report saved to: {output_path}[/green]")
    
    def _export_to_json(self, results: Dict, output_path: str):
        """Export to JSON format"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, indent=2, default=str)
        
        console.print(f"[green]JSON report saved to: {output_path}[/green]")
    
    def _export_to_markdown(self, results: Dict, output_path: str):
        """Export to Markdown format"""
        md_content = f"""# Document Comparison Report

Generated: {results['timestamp']}
Algorithm: {results['algorithm']}

## Files Compared

- **File 1:** {results['files']['file1']['path']}
- **File 2:** {results['files']['file2']['path']}

## Statistics

| Metric | File 1 | File 2 | Difference |
|--------|--------|--------|------------|
| Characters | {results['statistics']['text1']['characters']:,} | {results['statistics']['text2']['characters']:,} | {abs(results['statistics']['text1']['characters'] - results['statistics']['text2']['characters']):,} |
| Words | {results['statistics']['text1']['words']:,} | {results['statistics']['text2']['words']:,} | {abs(results['statistics']['text1']['words'] - results['statistics']['text2']['words']):,} |
| Lines | {results['statistics']['text1']['lines']:,} | {results['statistics']['text2']['lines']:,} | {abs(results['statistics']['text1']['lines'] - results['statistics']['text2']['lines']):,} |

## Similarity

**{results['statistics'].get('similarity_percentage', 0):.2f}%**

## Differences

```diff
"""
        
        # Add differences (limited)
        for diff in results['differences'][:100]:
            if diff.get('type') == 'addition':
                md_content += f"+ {diff.get('content', '')}\n"
            elif diff.get('type') == 'deletion':
                md_content += f"- {diff.get('content', '')}\n"
            elif diff.get('type') == 'context':
                md_content += f"  {diff.get('content', '')}\n"
        
        md_content += "```"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        # Report saved 